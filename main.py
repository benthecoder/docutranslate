import argparse
import base64
import os
from io import BytesIO

import fitz  # PyMuPDF
from docx import Document
from openai import OpenAI
from PIL import Image
from tqdm import tqdm

# Set the API key and model name
MODEL = "gpt-4o"
client = OpenAI(api_key="YOUR_API_KEY")


def pdf_page_to_base64(pdf_path, page_number):
    """Converts a PDF page to a base64 encoded image."""
    doc = fitz.open(pdf_path)
    page = doc.load_page(page_number)
    pix = page.get_pixmap()
    buffered = BytesIO()
    Image.frombytes("RGB", [pix.width, pix.height], pix.samples).save(
        buffered, format="PNG"
    )
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


def extract_and_translate_text(image_base64, model, target_language):
    """Extracts text from the image and translates it using OpenAI."""
    extract_prompt = {
        "role": "system",
        "content": "You are a helpful assistant.",
    }
    user_extract_message = {
        "role": "user",
        "content": [
            {
                "type": "text",
                "text": "Extract the text from the following image. Try to keep the same formatting you see in the image, such as tables and lists. Only return the extracted text in your response, don't include anything else.",
            },
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{image_base64}"},
            },
        ],
    }

    extract_response = client.chat.completions.create(
        model=model,
        messages=[extract_prompt, user_extract_message],
        temperature=0.0,
    )
    extracted_text = extract_response.choices[0].message.content.strip()

    translate_prompt = {
        "role": "system",
        "content": "You are an expert translator.",
    }
    user_translate_message = {
        "role": "user",
        "content": f"""
        Translate the following text to {target_language}. Follow these guidelines:

        1. **Translate all text.** Ensure the translation is entirely in {target_language} and does not contain any English.
        2. **Translate technical terms and jargon accurately.** Ensure that specialized terminology is translated appropriately within the context.
        3. **Retain the context and tone of the original text.** Make sure the translation is contextually appropriate and retains the original meaning and tone.
        4. **Verify the accuracy of the translation.** Double-check the translation to ensure it is free from errors and accurately represents the source text.
        \n\n{extracted_text}""",
    }

    translate_response = client.chat.completions.create(
        model=model,
        messages=[translate_prompt, user_translate_message],
        temperature=0.0,
    )
    translated_text = translate_response.choices[0].message.content.strip()

    return extracted_text, translated_text


def process_pdf(pdf_path, model, target_language):
    """Processes the entire PDF and returns original texts and translations."""
    doc = fitz.open(pdf_path)
    num_pages = doc.page_count
    original_texts = []
    translations = []

    print("Processing PDF...")
    for page_number in tqdm(range(num_pages), desc="Pages"):
        base64_image = pdf_page_to_base64(pdf_path, page_number)
        extracted_text, translated_text = extract_and_translate_text(
            base64_image, model, target_language
        )
        original_texts.append(extracted_text)
        translations.append(translated_text)
        print(
            f"Page {page_number + 1}: Extracted {len(extracted_text)} characters, Translated {len(translated_text)} characters"
        )

    return original_texts, translations


def create_word_doc(texts, output_path, title):
    """Creates a Word document with the given texts."""
    doc = Document()
    for i, text in enumerate(texts):
        doc.add_heading(f"Page {i + 1}", level=1)
        doc.add_paragraph(text)
    doc.save(output_path)
    print(f"{title} saved to {output_path}")


def process_single_page(pdf_path, page_number, model, target_language, output_txt_path):
    """Processes a single page of the PDF and saves the result to a text file."""
    base64_image = pdf_page_to_base64(pdf_path, page_number)
    extracted_text, translated_text = extract_and_translate_text(
        base64_image, model, target_language
    )
    with open(output_txt_path, "w", encoding="utf-8") as f:
        f.write(
            f"Original Text:\n{extracted_text}\n\nTranslated Text:\n{translated_text}"
        )
    print(
        f"Page {page_number + 1}: Extracted {len(extracted_text)} characters, Translated {len(translated_text)} characters"
    )
    print(f"Single page processing saved to {output_txt_path}")


def main():
    parser = argparse.ArgumentParser(description="Process and translate PDF documents.")
    parser.add_argument("pdf_path", type=str, help="Path to the PDF file to process.")
    parser.add_argument(
        "--single-page", action="store_true", help="Process a single page of the PDF."
    )
    parser.add_argument(
        "--page-number",
        type=int,
        default=1,
        help="Page number to process in single page mode.",
    )
    parser.add_argument(
        "--language",
        type=str,
        help="Target language for translation.",
    )

    args = parser.parse_args()

    pdf_path = args.pdf_path
    output_prefix = os.path.splitext(os.path.basename(pdf_path))[0]
    target_language = args.language

    original_texts_path = f"{output_prefix}_original_text.docx"
    translations_path = f"{output_prefix}_translated_{target_language}.docx"
    output_txt_path = f"{output_prefix}_page{args.page_number}_translation.txt"

    if args.single_page:
        process_single_page(
            pdf_path, args.page_number - 1, MODEL, target_language, output_txt_path
        )
    else:
        original_texts, translations = process_pdf(pdf_path, MODEL, target_language)
        create_word_doc(original_texts, original_texts_path, "Original texts")
        create_word_doc(translations, translations_path, "Translations")


if __name__ == "__main__":
    main()
